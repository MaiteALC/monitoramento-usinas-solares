""" Este módulo contém as funções para criação e formatação dos arquivos docx onde os prints diários do monitoramento são inseridos.
Tambem inclui a função que organiza e a que insere os prints em seus respectivos arquivos."""

import docx
from docx.document import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import locale
import logging
from monitoramento import enviar_email
from config import *


logger = logging.getLogger('Organização dos prints')

logger.setLevel(logging.INFO)

file_handler = logging.FileHandler(Path(CAMINHO_PASTA_LOGS, 'organizacao_prints.log'), mode='a', encoding='utf-8')

file_formatter = logging.Formatter(FORMATACAO_LOGGING)
file_handler.setFormatter(file_formatter)

logger.addHandler(file_handler)
logger.addHandler(logging.StreamHandler)


locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')



def criar_docx_monitoramentos(nome_usina: str, site: str) -> Document:
    """ Cria um novo docx, formatando a página incial.
    
    A função cria um novo arquivo docx, formata a primeira página com título, cabeçalho, rodapé, logo Apollo e informações do documento.
    O cabeçalho e rodapé são adicionados na primeira seção do documento, desde que mais nenhuma outra seja adicionada eles continuarão sendo herdados nas próximas páginas do arquivo.

    Args:
        nome_usina (str): o nome da usina que será adicionado na primeira página

    Returns:
        Document: o objeto do documento recém criado e formatado.

    """
    try:
        novo_doc = docx.Document()

        section = novo_doc.sections[0]

        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        header = section.header
        footer = section.footer

        header.paragraphs[0].text = ''
        footer.paragraphs[0].text = ''

        paragraph_header = header.paragraphs[0]
        paragraph_footer = footer.paragraphs[0]

        paragraph_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_header = paragraph_header.add_run()
        run_footer = paragraph_footer.add_run()

        run_header.add_picture(Path(CAMINHO_PASTA_RAIZ, 'img', 'cabecalho.png'), width=Cm(16.43), height=Cm(0.29)) 
        run_footer.add_picture(Path(CAMINHO_PASTA_RAIZ,'img', 'cabecalho.png'), width=Cm(16.43), height=Cm(0.29)) 

        novo_doc.add_picture(Path(CAMINHO_PASTA_RAIZ, 'logo_apollo.jpg'), width=Cm(17.1), height=Cm(8.5))

        ultimo_paragrafo = novo_doc.paragraphs[-1]
        ultimo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        titulo = novo_doc.add_paragraph()
        titulo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_titulo = titulo.add_run('Prints do relatório de monitoramento\n')

        run_titulo.font.name = 'Calibri'
        run_titulo.font.size = Pt(42)
        run_titulo.font.bold = True

        infos = novo_doc.add_paragraph()
        run_nome_usina = infos.add_run(f'\n\n\n\n\n\n{nome_usina}\n')

        mes_atual = DATA_ATUAL.strftime('%B')
        run_data_local = infos.add_run(f'ITAÚNA/MG\n{mes_atual} {DATA_ATUAL.year}')

        run_nome_usina.bold = True
        run_nome_usina.font.name = 'Calibri'
        run_nome_usina.font.size = Pt(18)

        run_data_local.font.name = 'Arial'
        run_data_local.font.size = Pt(14)

        novo_doc.save(Path(CAMINHO_PASTA_RAIZ, 'Histórico de monitoramentos', site, f'{nome_usina} - mês {DATA_ATUAL.month}.docx'))

    except Exception as e:
        logger.error(f'Erro {e} durante a criação do arquivo docx para a usina {site} - {nome_usina}')
        return

    else:
        logger.info(f'Criado arquivo docx para a usina {site} - {nome_usina}')
        return novo_doc



def organizar_screenshots(relacao_site_usina: dict) -> dict:
    """ Organiza screenshots, montando seus respectivos caminhos e os armazenando em um dicionário.

    Para cada usina de cada site que estiver na relacao_usina_site será feita uma lista com os caminhos para todas as screenshots do monitoramento daquela usina baseando-se no padrão de prints que são tirados para aquele site específico.

    Args:
        relacao_site_usina: um dicionário cujas chaves sejam os sites monitorados e os valores sejam as usinas daquele site.

    Returns:
        dict: Retorna um dicionário no formato:
            {
                'Site1': {
                    'UsinaA': [print1, print2, ...],
                    'UsinaB': [print1, print2, ...],
                },
                'Site2': { ...
                }
            }

    """
    screenshots = {}
    logger.info('Iniciando organização das screenshots')

    for site, usinas in relacao_site_usina.items():
        print(f'Organizando as screenshots do site {site}...')

        screenshots[site] = {}

        if site == 'PHB':
            for usina in usinas:
                try:
                    lista_prints = [
                        Path(CAMINHO_PASTA_PRINTS, 'PHB', f'{usina} - visão geral.png'),
                        *[Path(CAMINHO_PASTA_PRINTS, 'PHB', f'{usina} - inversor {n}.png') for n in range(1, 5)]
                    ]
                    screenshots[site][usina] = lista_prints 
                # Os caminhos para os prints dos inversores das usinas PHB são montados através de list comprehension devido a forma como os inversores são exibidos no site (em carrossel de imagens).

                except Exception as e:
                    logger.error(f'Erro ao organizar as screenshots da usina PHB - {usina}: {e}')
                    enviar_email('erro_no_codigo', site, usina, erro_capturado=e, onde_ocorreu_erro=f'organização dos prints da usina {usina} (PHB)')
                    continue   

        elif site == 'Sungrow':
            for usina in usinas:
                try:
                    lista_prints = [
                        Path(CAMINHO_PASTA_PRINTS, 'Sungrow', f'{usina} - visão geral.png'),
                        Path(CAMINHO_PASTA_PRINTS, 'Sungrow', f'{usina} - gráfico.png'),
                        Path(CAMINHO_PASTA_PRINTS, 'Sungrow', f'{usina} - inversores.png')

                    ]
                    screenshots[site][usina] = lista_prints 

                except Exception as e:
                    logger.error(f'Erro ao organizar as screenshots da usina Sungrow {usina}: {e}')
                    enviar_email('erro_no_codigo', site, usina, erro_capturado=e, onde_ocorreu_erro=f'organização dos prints da usina {usina} (Sungrow)')
                    continue   

        else:
            for usina in usinas:
                try:
                    lista_prints = [
                        Path(CAMINHO_PASTA_PRINTS, site, f'{usina} - visão geral.png'),
                        Path(CAMINHO_PASTA_PRINTS, site, f'{usina} - inversores.png')
                    ]
                    screenshots[site][usina] = lista_prints

                except Exception as e:
                    logger.error(f'Erro ao organizar as screenshots da usina {site} {usina}: {e}')
                    enviar_email('erro_no_codigo', site, usina, erro_capturado=e, onde_ocorreu_erro=f'organização dos prints da usina {usina} ({site})')
                    continue

        print(f'Organização das screenshots do site {site} concluída\n')

    logger.info('Organização das screenshots concluída.')

    return screenshots



def inserir_prints_docx(relacao_site_usina: dict, screenshots_organizadas: dict):
    """ Insere os prints do monitoramento no seu respectivo arquivo docx.
    
    Essa função irá inserir o print fornecido como argumento no docx daquela usina naquele mês específico, adicionando uma nova página em cada chamada e inserindo os prints na ordem em que foram fornecidos na lista.
    O ajuste de tamanho para as imagens é feito na própria função baseado no tamanho dos prints que foram tirados no mês 07/2025.

    Args:
        relacao_site_usina (dict): um dicionário que contenha os sites como chave e uma lista com as usinas de cada site como valor.

    Raises:
        FileNotFoundError: essa exceção será levantada caso o arquivo do docx ou o caminho para a imagem não for encontrado.

    """
    logger.info('Iniciando inserção dos prints nos respectivos arquivos docx')

    for site in relacao_site_usina.keys():
        usinas = relacao_site_usina[site]

        for nome_usina in usinas:
            caminho_doc = Path(CAMINHO_PASTA_DOCX, site, f'{nome_usina} - mês {DATA_ATUAL.month}.docx')

            if caminho_doc.exists:
                doc = docx.Document(caminho_doc)

            elif DATA_ATUAL.day == 1:
                doc = criar_docx_monitoramentos(nome_usina, site)

            else:
                logger.error(f'Arquivo docx do monitoramento para a usina {nome_usina} não encontrado. Continuando para o próximo...')
                continue

            nova_section = doc.add_section()

            nova_section.bottom_margin = Cm(2.5)
            nova_section.top_margin = Cm(2.5)

            nova_section.right_margin = Cm(0.4)
            nova_section.left_margin = Cm(0.4)

            data_str = AGORA.strftime('%d/%m/%Y')

            data_cabecalho = doc.add_paragraph()
            run_data = data_cabecalho.add_run(data_str)

            data_cabecalho.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run_data.font.name = 'Calibri'
            run_data.font.size = Pt(20)
            run_data.font.bold = True

            try:
                for screenshot in screenshots_organizadas[site][nome_usina]:
                    if site == 'Shine':
                        doc.add_picture(screenshot, width=Cm(18.3), height=Cm(9.5))

                        ultimo_paragrafo = doc.paragraphs[-1]
                        ultimo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    elif site == 'Sungrow':
                        doc.add_picture(screenshot, width=Cm(18))

                        ultimo_paragrafo = doc.paragraphs[-1]
                        ultimo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    else:
                        doc.add_picture(screenshot, width=Cm(20))

                        ultimo_paragrafo = doc.paragraphs[-1]
                        ultimo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

            except Exception as e:
                logger.error(f'Erro ao inserir as screenshots da usina {nome_usina} no docx: {e}')

            else:
                doc.save(caminho_doc)
                logger.info(f'Prints da usina {nome_usina} inserido no docx com sucesso!')

    logger.info('Inserção dos prints nos arquivos docx concluído com sucesso!') 
